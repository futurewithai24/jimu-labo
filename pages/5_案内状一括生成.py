import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io
import zipfile

from utils import show_logo
show_logo()
st.page_link("home.py", label="← ホームに戻る")
st.title("📝 案内状一括生成")
st.caption("送付先リストのExcelをアップするだけで、人数分のWord案内状を自動作成します")

# ── サイドバー（設定） ──────────────────────────
with st.sidebar:
    st.header("文書の設定")
    title_text  = st.text_input("文書タイトル",  value="研修のご案内")
    sender_name = st.text_input("差出人名",       value="総務部")
    honorific   = st.text_input("敬称",           value="様")
    st.divider()
    st.header("Excelの列名設定")
    col_name  = st.text_input("氏名の列名",   value="氏名")
    col_dept  = st.text_input("部署の列名",   value="部署")
    col_date  = st.text_input("研修日の列名", value="研修日")
    col_venue = st.text_input("会場の列名",   value="会場")

# ── サンプルExcelダウンロード ──────────────────
sample_df = pd.DataFrame({
    "氏名":  ["田中 花子", "佐藤 美咲", "山田 恵子"],
    "部署":  ["総務部", "営業部", "経理部"],
    "研修日": ["2026/07/10", "2026/07/10", "2026/07/11"],
    "会場":  ["本社 会議室A", "本社 会議室A", "本社 会議室B"],
})
sample_buf = io.BytesIO()
sample_df.to_excel(sample_buf, index=False)
sample_buf.seek(0)
st.download_button(
    label="サンプルExcelをダウンロード",
    data=sample_buf,
    file_name="送付先リスト_サンプル.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

st.divider()

# ── Excelアップロード ──────────────────────────
uploaded_file = st.file_uploader(
    "送付先リストのExcelをアップロード",
    type=["xlsx", "xls"]
)

def format_date(value) -> str:
    if hasattr(value, "strftime"):
        return value.strftime("%Y年%m月%d日")
    return str(value)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_document(name, dept, date, venue, today, title_text, sender_name, honorific):
    doc = Document()

    # 日付（右寄せ）
    date_para = doc.add_paragraph(today)
    date_para.alignment = 2

    # 宛名
    doc.add_paragraph(f"{dept}　{name} {honorific}")
    doc.add_paragraph("")

    # 差出人（右寄せ）
    sender_para = doc.add_paragraph(sender_name)
    sender_para.alignment = 2

    # タイトル（中央・太字）
    title_para = doc.add_paragraph()
    title_para.alignment = 1
    run = title_para.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph("拝啓　時下ますますご清栄のこととお慶び申し上げます。")
    doc.add_paragraph("下記のとおり研修を実施いたしますので、ご参加くださいますようご案内申し上げます。")
    doc.add_paragraph("")

    # テーブル（罫線なし・列幅広め）
    table = doc.add_table(rows=3, cols=2)
    remove_table_borders(table)
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(13.0)
    table.cell(0, 0).text = "日時"
    table.cell(0, 1).text = date
    table.cell(1, 0).text = "会場"
    table.cell(1, 1).text = venue
    table.cell(2, 0).text = "対象"
    table.cell(2, 1).text = f"{dept} 全員"

    doc.add_paragraph("")

    # 敬具（右寄せ）
    keigo_para = doc.add_paragraph("敬具")
    keigo_para.alignment = 2

    return doc

# ── メイン処理 ──────────────────────────────────
if uploaded_file is not None:
    try:
        df = pd.read_excel(uploaded_file)
        st.success(f"{len(df)} 件のデータを読み込みました")
        st.dataframe(df, use_container_width=True)

        if st.button("案内状を一括生成する", type="primary"):
            with st.spinner("作成中です。しばらくお待ちください…"):
                today   = datetime.today().strftime("%Y年%m月%d日")
                results = []
                zip_buffer = io.BytesIO()

                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, row in df.iterrows():
                        try:
                            name  = str(row[col_name])
                            dept  = str(row[col_dept])
                            date  = format_date(row[col_date])
                            venue = str(row[col_venue])

                            doc      = create_document(name, dept, date, venue, today, title_text, sender_name, honorific)
                            filename = f"{dept}_{name}_案内状.docx"

                            doc_buf = io.BytesIO()
                            doc.save(doc_buf)
                            doc_buf.seek(0)
                            zf.writestr(filename, doc_buf.read())
                            results.append({"ファイル名": filename, "結果": "✅ 完了"})

                        except Exception as e:
                            results.append({"ファイル名": f"行{i+2}", "結果": f"❌ エラー：{e}"})

                zip_buffer.seek(0)

            st.success(f"完成しました！{len(df)} 件の案内状を作成しました")
            st.dataframe(pd.DataFrame(results), use_container_width=True)
            st.download_button(
                label="ZIPをダウンロード（Word一式）",
                data=zip_buffer,
                file_name="案内状一式.zip",
                mime="application/zip"
            )

    except KeyError as e:
        st.error(f"列名が見つかりません：{e}　左のサイドバーで列名を確認してください")
    except Exception as e:
        st.error(f"エラーが発生しました：{e}")

with st.expander("📖 使い方"):
    st.write("1. サンプルExcelをダウンロードして送付先を入力・保存")
    st.write("2. 左のサイドバーで文書タイトル・差出人・列名を設定")
    st.write("3. Excelをアップロードして「案内状を一括生成する」ボタンを押す")
    st.write("4. ZIPをダウンロードすると人数分のWordファイルが入っています")
    st.write("- 列名がExcelと違う場合は左サイドバーで変更できます")
